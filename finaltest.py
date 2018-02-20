from tkinter import *
from docx import Document
from docx.shared import Inches
from docx.shared import Pt

def f():
    global v1,v2,v3
    v1 = professor_name_entry.get()
    v2 = professor_desg_entry.get()
    v3 = professor_colg_entry.get()
def generate_doc():
    v1 = professor_name_entry.get()
    v2 = professor_desg_entry.get()
    v3 = professor_colg_entry.get()
    document = Document('NOC_blank.docx')

    p = document.add_paragraph()
    p.style = document.styles['Normal']
    font = p.style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    p.add_run('\nTo, \n').bold = True
    p.add_run(v1).bold=True
    p.add_run(',\n')
    p.add_run(v2).bold=True
    p.add_run(',')
    p.add_run('\nDepartment of Computer Science and Engineering,\n').bold = True
    p.add_run(v3).bold=True
    #p.style= document.styles['Body Text 1']

    #p.add_run('\nProf.Shashidhar G Koolagudi,').bold = True
    #p.add_run('\nAssistant Professor,').bold = True
    #p.add_run('\nDepartment of Computer Science and Engineering,').bold = True
    #p.add_run('\nNIT Surathkal').bold = True
    # paragraph_format = p.paragraph_format
    # paragraph_format.left_indent = Inches(0.5)

    r = document.add_paragraph()
    r.style = document.styles['Normal']
    font = r.style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    r.add_run('Sub: Summer Internship during Summer Vacation for CSE Students').bold = True
    paragraph_format = r.paragraph_format
    paragraph_format.left_indent = Inches(1)
    document.add_paragraph('Dear Sir,')

    s = document.add_paragraph('Greetings from')
    s.style = document.styles['Normal']
    font = s.style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    s.add_run(' Jalpaiguri Government Engineering College!').bold = True

    p = document.add_paragraph()
    p.style = document.styles['Normal']
    font = p.style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    p.add_run(' I am glad to inform you that our curriculum is designed for all 2nd year students to '
              ' undergo internship in academic ')
    p.add_run(' institutes for eight weeks. During this period starting from ')
    p.add_run(' 15th May to 15th July, 2018,').bold = True
    p.add_run(' they are expected to be involved in academic projects. ')
    p.add_run(' We would be grateful to you, if you could accommodate them at your esteemed organization:')
    p_format = p.paragraph_format
    p_format.left_indent = Inches(0.5)

    items = []

    for i in range(stud_count):
        s2 = stud_roll_entry_list[i].get()
        s1 = stud_name_entry_list[i].get()
        s3 = stud_email_entry_list[i].get()
        s4 = stud_phno_entry_list[i].get()

        obj = {
            "sl": i + 1,
            "roll_no": s2,
            "name": s1,
            "email": s3,
            'phone_no': s4
        }

        items.append(obj)

    '''items = [
        {
            "sl": 1,
            "roll_no": 16101104056,
            "name": "Anisha Datta",
            "email": 'ad2056@cse.jgec.ac.in',
            'phone_no': '07688090024'
        } 
    ]'''
    print(items)
    # add table ------------------
    col_names = ('Sl', 'Roll No', 'Name', 'Email', 'Phone No')
    table = document.add_table(rows=1, cols=len(col_names))
    heading_cells = table.rows[0].cells
    for idx, name in enumerate(col_names):
        paragraph = heading_cells[idx].paragraphs[0]
        run = paragraph.add_run(name)
        run.bold = True
    table.style = 'TableGrid'

    for item in items:
        cells = table.add_row().cells
        cells[0].text = str(item["sl"])
        cells[1].text = str(item["roll_no"])
        cells[2].text = str(item["name"])
        cells[3].text = str(item["email"])
        cells[4].text = str(item["phone_no"])

    q = document.add_paragraph()
    q.add_run(
        '\n\nWe eagerly look forward to hear from you soon. The department has no objection if the students undergo Summer ')
    q.add_run('Internship during 15th May 2017- 15thJuly, 2017.')
    q.add_run('\n\nThanking you.')
    q.add_run('\nYours sincerely,')
    q.add_run('\nDr. Dipak Kumar Kole')
    q.add_run('\n\n\n___________________________________')
    q.add_run('\nDr. Dipak Kumar Kole').bold=True
    q.add_run('\nHead Of the Dept.').bold=True
    q.add_run('\nComputer Science & Engineering Dept.').bold=True
    q.add_run('\nJalpaiguri Govt. Engineering College(Autonomous)').bold=True
    document.save('sample5.docx')


def get_stud_count():
    global student_count_entry
    cnt = student_count_entry.get()
    try:
        global stud_count
        cnt = int(cnt)
        stud_count = cnt
        students_widget_generation()
    except ValueError:
        student_count_entry.delete(0, 'end')


def students_widget_generation():
    global stud_name_entry_list
    global stud_roll_entry_list
    global stud_email_entry_list
    global stud_phno_entry_list

    stud_name_entry_list = []
    stud_roll_entry_list = []
    stud_email_entry_list = []
    stud_phno_entry_list = []

    stud_name_label_list = []
    stud_roll_label_list = []
    stud_email_label_list = []
    stud_phno_label_list = []
    stud_head_label_list = []

    for i in range(stud_count):
        stud_head_label_list.append(Label(main_window, text='Student #' + str(i + 1)))
        stud_name_label_list.append(Label(main_window, text='Name #' + str(i + 1)))
        stud_roll_label_list.append(Label(main_window, text='Roll #' + str(i + 1)))
        stud_email_label_list.append(Label(main_window, text='Email #' + str(i + 1)))
        stud_phno_label_list.append(Label(main_window, text='Mob #' + str(i + 1)))

    row_count = 7

    for i in range(stud_count):

        stud_head_label_list[i].grid(row=row_count, column=4 if (i % 2) else 2)
        row_count += 1

        stud_name_label_list[i].grid(row=row_count, column=3 if (i % 2) else 1)
        ent = Entry(main_window)
        ent.grid(row=row_count, column=4 if (i % 2) else 2)
        stud_name_entry_list.append(ent)
        row_count += 1

        stud_roll_label_list[i].grid(row=row_count, column=3 if (i % 2) else 1)
        ent = Entry(main_window)
        ent.grid(row=row_count, column=4 if (i % 2) else 2)
        stud_roll_entry_list.append(ent)
        row_count += 1

        stud_email_label_list[i].grid(row=row_count, column=3 if (i % 2) else 1)
        ent = Entry(main_window)
        ent.grid(row=row_count, column=4 if (i % 2) else 2)
        stud_email_entry_list.append(ent)
        row_count += 1

        stud_phno_label_list[i].grid(row=row_count, column=3 if (i % 2) else 1)
        ent = Entry(main_window)
        ent.grid(row=row_count, column=4 if (i % 2) else 2)
        stud_phno_entry_list.append(ent)
        row_count += 1

        print(row_count)
        if not i % 2:
            row_count -= 5

    btn2 = Button(main_window, text='Generate NOC', command=generate_doc)
    btn2.grid(row=row_count + 10, column=2)


main_window = Tk()
main_window.title("#NOC Generation#")
main_window.geometry('700x700+200+100')
professor_main_label = Label(main_window, text="DETAILS OF PROFESSOR")
professor_main_label.grid(row=1, column=2)

professor_name_label = Label(main_window, text='Name')
professor_name_label.grid(row=2, column=1)
professor_name_entry = Entry(main_window)
professor_name_entry.grid(row=2, column=2)


professor_desg_label = Label(main_window, text='Designation')
professor_desg_label.grid(row=3, column=1)
professor_desg_entry = Entry(main_window)
professor_desg_entry.grid(row=3, column=2)

professor_colg_label = Label(main_window, text='College')
professor_colg_label.grid(row=4, column=1)
professor_colg_entry = Entry(main_window)
professor_colg_entry.grid(row=4, column=2)

student_count_label = Label(main_window, text='No. of students')
student_count_label.grid(row=5, column=1)
student_count_entry = Entry(main_window)
student_count_entry.grid(row=5, column=2)

#v1= professor_name_entry.get()
#v2= professor_desg_entry.get()
#v3= professor_colg_entry.get()

btn1 = Button(main_window, text='Enter Student Details', command=get_stud_count)
btn1.grid(row=6, column=2)

main_window.mainloop()
